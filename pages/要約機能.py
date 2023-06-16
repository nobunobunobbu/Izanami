import streamlit as st
from PyPDF2 import PdfReader
import requests
import textwrap
from PIL import Image
import common
from datetime import datetime
import docx
from docx.shared import Pt
from pydub import AudioSegment
import os

image = Image.open('IZANAMI.png')

st.image(image,use_column_width=True)

common.check_login()

# Streamlit アプリケーションの設定
st.title("要約機能")


tab1, tab2 = st.tabs(["PDF 要約機能","音声文字起こし・要約機能"])

now = datetime.now()
date_str = now.strftime("%Y-%m-%d")


# ChatGPT のAPIキー入力用テキストボックス
api_key = st.secrets["api_key"]["api_key"]
st.session_state.api_key = api_key

with tab1:

 # ユーザーからPDF ファイルをアップロード
 uploaded_file = st.file_uploader("PDF ファイルをアップロードしてください", type="pdf")

 # モデル選択のセレクトボックス
 model = st.selectbox('GPTモデルの選択', ['gpt-3.5-turbo', 'gpt-3.5-turbo-16k', 'gpt-4'],help="gpt-3.5-turbo-16k :より長文のプロンプトを受け付けます。gpt-4:未対応となります。")

 # 実行ボタン
 run_button = st.button('実行')

 if run_button:
    if uploaded_file is not None and api_key != "":
        # アップロードされたPDF ファイルを解析してテキストを抽出
        pdf_reader = PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()

    # テキストをトークン数に合わせて分割
    chunks = textwrap.wrap(text, 2000)

    explanations = []

    for chunk in chunks:
        # ChatGPT にテキストを送信して解説を生成
        prompt = "テキストを読んで内容をマークダウンでわかりやすく解説して"
        headers = {
            "Authorization": f"Bearer {st.session_state.api_key}",
            "Content-Type": "application/json"
        }
        data = {
            "model": model,
            "messages": [
                {
                    "role": "system",
                    "content": "You are a helpful assistant."
                },
                {
                    "role": "user",
                    "content": prompt + chunk
                }
            ]
        }
        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=data)

        if response.status_code == 200:
            try:
                explanation = response.json()["choices"][0]["message"]["content"].strip()
                explanations.append(explanation)
            except KeyError:
                st.error("Failed to retrieve explanation from the API response.")
        else:
            try:
                error_details = response.json()
            except ValueError:
                error_details = "No additional information."
            st.error(f"API request failed with status code: {response.status_code}. Details: {error_details}")

    # All explanations are displayed as a single text
    st.subheader("▼")
    st.write(' '.join(explanations))
    doc = docx.Document()

    # テキストを追加
    doc.add_paragraph(' '.join(explanations))

    # ドキュメントを保存
    doc.save("answer.docx")

    # 保存したドキュメントを読み込み、ダウンロードボタンのデータとして指定
    with open('answer.docx', 'rb') as f:
     st.download_button(
        label="ダウンロード",
        data=f.read(),
        file_name=date_str +'_PDF要約.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
     
with tab2:
   option = st.selectbox("機能の選択", ["動画から音声を抜き出す", "音声文字起こし・要約"])

   if option == "動画から音声を抜き出す":
    # Upload video file
    uploaded_file = st.file_uploader("Upload a video file", type=["mp4", "mov", "avi"])

    if uploaded_file is not None:
        if st.button('実行',key="botton"):
            # Save the uploaded video file locally
            with open('temp_video.mp4', 'wb') as f:
                f.write(uploaded_file.getbuffer())

            # Convert video to audio using pydub
            video = AudioSegment.from_file('temp_video.mp4')
            video.export("temp_audio.wav", format="wav")

            # Download button for wav file
            with open("temp_audio.wav", "rb") as f:
                bytes = f.read()
            st.download_button(
                label="Download audio",
                data=bytes,
                file_name="audio.wav",
                mime="audio/wav",
            )

            # Clean up temporary files
            os.remove("temp_video.mp4")
            os.remove("temp_audio.wav")


   elif option == "音声文字起こし・要約":
     uploaded_file = st.file_uploader("Upload an audio file", type=["wav", "mp3","m4a"])
     model = st.selectbox('GPTモデルの選択', ['gpt-3.5-turbo','gpt-3.5-turbo-16k', 'gpt-4'],help="gpt-3.5-turbo-0613 :より長文のプロンプトを受け付けます。gpt-4:")


     if uploaded_file is not None:
         if st.button('実行',key="botton2"):
            # Load the uploaded file into memory
            audio_data = uploaded_file.read()

            # Create a request for the OpenAI's Whisper API
            url = "https://api.openai.com/v1/audio/transcriptions"
            headers = {
                "Authorization": f"Bearer {api_key}",  # Replace with your OpenAI API key
            }
            
            # Prepare multipart encoded data
            files = {
                "file": ("audio.mp3", audio_data, "audio/mpeg"),  # Modify as needed
            }
            
            # Model data
            data = {
                "model": "whisper-1"
            }

            # Send a POST request and get the response
            response = requests.post(url, headers=headers, files=files, data=data)

            if response.status_code == 200:
                transcription = response.json()['text']
                with st.expander("文字起こし結果"):
                        st.write(transcription)

                # Use the chat-based GPT-3 API to summarize the transcription
                headers = {
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                }
                data = {
                    "model": model,
                    "messages": [
                        {
                            "role": "system",
                            "content": "You are a helpful assistant."
                        },
                        {
                            "role": "user",
                            "content": f"以下の文章を分かりやすくマークダウン形式で要約して: {transcription}"
                        }
                    ]
                }

                response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=data)
                if response.status_code == 200:
                    response_data = response.json()
                    summarized_text = response_data['choices'][0]['message']['content']
                    st.subheader("要約：")
                    st.write( summarized_text)
                else:
                    st.write("Error in summarization. Details:", response.json())

            else:
                st.write("Error in transcription. Details:", response.json())

            doc = docx.Document()
    # テキストを追加
            doc.add_paragraph(' '.join(summarized_text))

    # ドキュメントを保存
            doc.save("answer.docx")

    # 保存したドキュメントを読み込み、ダウンロードボタンのデータとして指定
            with open('answer.docx', 'rb') as f:
             st.download_button(
        label="ダウンロード",
        data=f.read(),
        file_name=date_str +'_文字起こし要約.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


